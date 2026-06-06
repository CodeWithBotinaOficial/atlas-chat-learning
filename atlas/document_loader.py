"""
Load and extract text from local or remote documents (PDF, DOCX, TXT, MD).
"""

import io
import os
import re
import tempfile
from urllib.parse import urlparse

import requests

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}
CONTENT_TYPE_TO_EXTENSION = {
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
    'text/plain': '.txt',
    'text/markdown': '.md',
}


class DocumentLoadError(Exception):
    """Raised when a document cannot be loaded or parsed."""


def sanitize_text(text):
    """
    Normalize whitespace and line breaks in extracted text.
    """
    if not text:
        return ""

    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)

    cleaned_lines = []
    for line in text.split('\n'):
        line = line.strip()
        if line:
            cleaned_lines.append(line)

    return '\n\n'.join(cleaned_lines).strip()


def chunk_text(text, max_chunk_size=1000):
    """
    Split text into chunks respecting word boundaries.
    """
    if not text or not text.strip():
        return []

    words = text.split()
    chunks = []
    current_chunk = []
    current_chunk_len = 0

    for word in words:
        word_len = len(word)
        separator_len = 1 if current_chunk else 0

        if current_chunk and current_chunk_len + separator_len + word_len > max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_chunk_len = word_len
        else:
            current_chunk.append(word)
            current_chunk_len += separator_len + word_len

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks


def detect_format_from_extension(path):
    """
    Detect supported document format from a file path or URL path.
    """
    _, ext = os.path.splitext(path)
    ext = ext.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentLoadError(
            f"Unsupported file format '{ext or '(none)'}'. "
            f"Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return ext


def detect_format_from_content_type(content_type):
    """
    Detect supported document format from an HTTP Content-Type header.
    """
    if not content_type:
        return None

    mime_type = content_type.split(';')[0].strip().lower()
    return CONTENT_TYPE_TO_EXTENSION.get(mime_type)


def extract_text_from_pdf(file_path_or_bytes):
    """
    Extract text from a PDF file path or bytes object.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentLoadError(
            "Missing dependency 'pypdf'. Install it with: pip install pypdf"
        ) from exc

    try:
        if isinstance(file_path_or_bytes, (bytes, bytearray)):
            reader = PdfReader(io.BytesIO(file_path_or_bytes))
        else:
            reader = PdfReader(file_path_or_bytes)

        page_texts = []
        for page in reader.pages:
            page_text = page.extract_text() or ''
            if page_text.strip():
                page_texts.append(page_text)

        if not page_texts:
            raise DocumentLoadError("No extractable text found in PDF.")

        return sanitize_text('\n\n'.join(page_texts))
    except DocumentLoadError:
        raise
    except Exception as exc:
        raise DocumentLoadError(f"Failed to extract text from PDF: {exc}") from exc


def extract_text_from_docx(file_path_or_bytes):
    """
    Extract text from a DOCX file path or bytes object.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentLoadError(
            "Missing dependency 'python-docx'. Install it with: pip install python-docx"
        ) from exc

    try:
        if isinstance(file_path_or_bytes, (bytes, bytearray)):
            document = Document(io.BytesIO(file_path_or_bytes))
        else:
            document = Document(file_path_or_bytes)

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        if not paragraphs:
            raise DocumentLoadError("No extractable text found in DOCX.")

        return sanitize_text('\n\n'.join(paragraphs))
    except DocumentLoadError:
        raise
    except Exception as exc:
        raise DocumentLoadError(f"Failed to extract text from DOCX: {exc}") from exc


def extract_text_from_txt(file_path):
    """
    Read plain text from a TXT or MD file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as file:
            text = file.read()
    except OSError as exc:
        raise DocumentLoadError(f"Failed to read text file: {exc}") from exc

    if not text.strip():
        raise DocumentLoadError("Text file is empty.")

    return sanitize_text(text)


def _extract_text_by_extension(file_path_or_bytes, extension):
    """
    Dispatch extraction to the appropriate handler for a supported extension.
    """
    if extension == '.pdf':
        return extract_text_from_pdf(file_path_or_bytes)
    if extension == '.docx':
        return extract_text_from_docx(file_path_or_bytes)
    if extension in {'.txt', '.md'}:
        if isinstance(file_path_or_bytes, (bytes, bytearray)):
            try:
                text = file_path_or_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = file_path_or_bytes.decode('latin-1')
            if not text.strip():
                raise DocumentLoadError("Text file is empty.")
            return sanitize_text(text)
        return extract_text_from_txt(file_path_or_bytes)

    raise DocumentLoadError(f"Unsupported file format '{extension}'.")


def load_from_local(path):
    """
    Load and extract cleaned text from a local document file.
    """
    if not path:
        raise DocumentLoadError("No file path provided.")

    resolved_path = os.path.abspath(os.path.expanduser(path))
    if not os.path.isfile(resolved_path):
        raise DocumentLoadError(f"File not found: {resolved_path}")

    print(f"[✓] Loading local document: {resolved_path}")
    extension = detect_format_from_extension(resolved_path)
    print(f"[✓] Detected format: {extension}")

    print("[✓] Extracting text...")
    text = _extract_text_by_extension(resolved_path, extension)
    print(f"[✓] Extracted {len(text)} characters.")
    return text


def load_from_url(url, temp_dir=None):
    """
    Download a remote document and extract cleaned text.
    """
    if not url:
        raise DocumentLoadError("No URL provided.")

    parsed_url = urlparse(url)
    if parsed_url.scheme not in {'http', 'https'}:
        raise DocumentLoadError("URL must use http or https.")

    print(f"[✓] Downloading document from: {url}")

    try:
        response = requests.get(url, stream=True, timeout=30)
        response.raise_for_status()
    except requests.exceptions.RequestException as exc:
        raise DocumentLoadError(f"Network error while downloading document: {exc}") from exc

    content_type = response.headers.get('Content-Type', '')
    extension = None
    if os.path.splitext(parsed_url.path)[1]:
        try:
            extension = detect_format_from_extension(parsed_url.path)
        except DocumentLoadError:
            extension = None

    if extension is None:
        extension = detect_format_from_content_type(content_type)

    if extension is None:
        raise DocumentLoadError(
            "Could not determine document format from URL extension or Content-Type. "
            f"Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    print(f"[✓] Detected format: {extension}")

    file_bytes = response.content
    if not file_bytes:
        raise DocumentLoadError("Downloaded document is empty.")

    temp_path = None
    try:
        if extension in {'.txt', '.md'}:
            print("[✓] Extracting text...")
            text = _extract_text_by_extension(file_bytes, extension)
        else:
            with tempfile.NamedTemporaryFile(delete=False, suffix=extension, dir=temp_dir) as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name

            print("[✓] Extracting text...")
            text = _extract_text_by_extension(temp_path, extension)

        print(f"[✓] Extracted {len(text)} characters.")
        return text
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
